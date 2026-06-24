import base64
from io import BytesIO
from zipfile import ZipFile

from docx import Document


def test_file_upload_list_analyze_and_delete_flow(client):
    upload_response = client.post(
        "/api/files/upload",
        data={"user_id": "browser-files-a", "category": "texto"},
        files={
            "file": (
                "erro-websocket.txt",
                b"Erro WebSocket no Render. Verifique wss e window.location.host.",
                "text/plain",
            )
        },
    )

    assert upload_response.status_code == 200
    uploaded = upload_response.json()["file"]
    assert uploaded["user_id"] == "browser-files-a"
    assert uploaded["category"] == "texto"
    assert uploaded["analysis_status"] == "pending"

    list_response = client.get("/api/files", params={"user_id": "browser-files-a"})
    assert list_response.status_code == 200
    assert [record["id"] for record in list_response.json()["files"]] == [uploaded["id"]]

    wrong_user_response = client.get(f"/api/files/{uploaded['id']}", params={"user_id": "browser-files-b"})
    assert wrong_user_response.status_code == 404

    analyze_response = client.post(
        f"/api/files/{uploaded['id']}/analyze",
        json={"user_id": "browser-files-a", "instructions": "explique o erro"},
    )

    assert analyze_response.status_code == 200
    analyzed = analyze_response.json()
    assert analyzed["file"]["analysis_status"] == "ready"
    assert "erro-websocket.txt" in analyzed["summary"]
    assert "websocket" in analyzed["keywords"]
    assert "pesquisar fontes atuais" in analyzed["message"]

    delete_response = client.delete(f"/api/files/{uploaded['id']}", params={"user_id": "browser-files-a"})
    assert delete_response.status_code == 200
    assert client.get(f"/api/files/{uploaded['id']}", params={"user_id": "browser-files-a"}).status_code == 404


def test_file_upload_blocks_dangerous_extension(client):
    response = client.post(
        "/api/files/upload",
        data={"user_id": "browser-files-safe", "category": "geral"},
        files={"file": ("setup.exe", b"MZ dangerous", "application/octet-stream")},
    )

    assert response.status_code == 400
    assert "bloqueado" in response.json()["detail"]


def test_camera_photo_upload_and_basic_analysis(client):
    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADElEQVR42mP8z8AARQAElwG5Nq7xWQAAAABJRU5ErkJggg=="
    )
    payload = {
        "user_id": "browser-camera-a",
        "filename": "camera.png",
        "image_data": f"data:image/png;base64,{base64.b64encode(png_bytes).decode()}",
    }

    upload_response = client.post("/api/camera/photo", json=payload)

    assert upload_response.status_code == 200
    uploaded = upload_response.json()["file"]
    assert uploaded["source"] == "camera"
    assert uploaded["category"] == "camera"

    analyze_response = client.post(
        f"/api/files/{uploaded['id']}/analyze",
        json={"user_id": "browser-camera-a", "manual_description": "print de erro do navegador"},
    )

    assert analyze_response.status_code == 200
    body = analyze_response.json()
    assert body["file"]["analysis_status"] == "limited"
    assert "OCR" in body["summary"]
    assert "print de erro do navegador" in body["summary"]


def test_document_uploads_extract_docx_xlsx_and_pptx_text(client):
    samples = [
        (
            "aula-orion.docx",
            build_docx_bytes("Orion explica matematica e programacao em uma apostila organizada."),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "matematica",
        ),
        (
            "planejamento.xlsx",
            build_xlsx_bytes("Projeto Orion", "Render 24/7", "WebSocket estavel"),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "render",
        ),
        (
            "slides.pptx",
            build_pptx_bytes("Professor Orion", "Flashcards e simulados"),
            "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            "flashcards",
        ),
    ]

    for filename, content, content_type, expected_keyword in samples:
        upload_response = client.post(
            "/api/files/upload",
            data={"user_id": "browser-docs-a", "category": "documento"},
            files={"file": (filename, content, content_type)},
        )
        assert upload_response.status_code == 200
        uploaded = upload_response.json()["file"]

        analyze_response = client.post(
            f"/api/files/{uploaded['id']}/analyze",
            json={"user_id": "browser-docs-a", "instructions": "resuma para estudar"},
        )

        assert analyze_response.status_code == 200
        body = analyze_response.json()
        assert body["file"]["analysis_status"] == "ready"
        assert expected_keyword in body["summary"].lower()


def test_file_transform_generates_pdf_and_download(client):
    upload_response = client.post(
        "/api/files/upload",
        data={"user_id": "browser-transform-a", "category": "texto"},
        files={
            "file": (
                "conteudo.txt",
                b"Orion pode transformar documentos em PDF, apostilas, trabalhos e flashcards para estudo.",
                "text/plain",
            )
        },
    )
    assert upload_response.status_code == 200
    source = upload_response.json()["file"]

    transform_response = client.post(
        f"/api/files/{source['id']}/transform",
        json={"user_id": "browser-transform-a", "mode": "pdf", "output_format": "pdf"},
    )

    assert transform_response.status_code == 200
    transformed = transform_response.json()
    assert transformed["generated_file"]["extension"] == ".pdf"
    assert transformed["generated_file"]["source"] == "generated"
    assert "Meus Arquivos" in transformed["message"]

    download_response = client.get(
        f"/api/files/{transformed['generated_file']['id']}/download",
        params={"user_id": "browser-transform-a"},
    )
    assert download_response.status_code == 200
    assert download_response.content.startswith(b"%PDF")


def test_file_transform_flashcards_creates_generated_text_file(client):
    upload_response = client.post(
        "/api/files/upload",
        data={"user_id": "browser-flashcards-a", "category": "texto"},
        files={
            "file": (
                "historia.txt",
                b"Memoria de longo prazo ajuda o Orion a lembrar projetos. Flashcards ajudam na revisao espacada.",
                "text/plain",
            )
        },
    )
    source = upload_response.json()["file"]

    transform_response = client.post(
        f"/api/files/{source['id']}/transform",
        json={"user_id": "browser-flashcards-a", "mode": "flashcards"},
    )

    assert transform_response.status_code == 200
    payload = transform_response.json()
    assert "Pergunta:" in payload["content"]
    assert payload["generated_file"]["extension"] == ".txt"


def build_docx_bytes(text: str) -> bytes:
    document = Document()
    document.add_heading("Documento Orion", level=1)
    document.add_paragraph(text)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def build_xlsx_bytes(*values: str) -> bytes:
    stream = BytesIO()
    with ZipFile(stream, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types xmlns='http://schemas.openxmlformats.org/package/2006/content-types'/>")
        archive.writestr("xl/sharedStrings.xml", "".join([
            "<sst xmlns='http://schemas.openxmlformats.org/spreadsheetml/2006/main'>",
            *[f"<si><t>{value}</t></si>" for value in values],
            "</sst>",
        ]))
        cells = "".join(
            f"<c r='A{index}' t='s'><v>{index - 1}</v></c>"
            for index, _value in enumerate(values, start=1)
        )
        archive.writestr(
            "xl/worksheets/sheet1.xml",
            "<worksheet xmlns='http://schemas.openxmlformats.org/spreadsheetml/2006/main'>"
            f"<sheetData><row r='1'>{cells}</row></sheetData></worksheet>",
        )
    return stream.getvalue()


def build_pptx_bytes(*texts: str) -> bytes:
    stream = BytesIO()
    with ZipFile(stream, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types xmlns='http://schemas.openxmlformats.org/package/2006/content-types'/>")
        text_nodes = "".join(f"<a:t>{text}</a:t>" for text in texts)
        archive.writestr(
            "ppt/slides/slide1.xml",
            "<p:sld xmlns:p='http://schemas.openxmlformats.org/presentationml/2006/main' "
            "xmlns:a='http://schemas.openxmlformats.org/drawingml/2006/main'>"
            f"<p:cSld><p:spTree>{text_nodes}</p:spTree></p:cSld></p:sld>",
        )
    return stream.getvalue()
