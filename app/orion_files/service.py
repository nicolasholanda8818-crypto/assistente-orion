from __future__ import annotations

import base64
import binascii
import json
import logging
import mimetypes
import re
import unicodedata
import zipfile
from collections import Counter
from io import BytesIO
from pathlib import Path
from typing import Literal
from uuid import uuid4
from xml.etree import ElementTree

from app.core.config import settings
from app.db import repositories
from app.orion_files.models import (
    OrionFileAnalysisResponse,
    OrionFileRecord,
    OrionFilesStatus,
    OrionFileTransformResponse,
    OrionFileUploadResponse,
)

LOGGER = logging.getLogger(__name__)
USER_ID_PATTERN = re.compile(r"^[A-Za-z0-9_.-]{1,64}$")
SAFE_NAME_PATTERN = re.compile(r"[^a-zA-Z0-9_.-]+")
TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json"}
IMAGE_EXTENSIONS = {".bmp", ".gif", ".jpeg", ".jpg", ".png", ".webp"}
DOCUMENT_EXTENSIONS = {".doc", ".docx", ".pdf", ".pptx", ".xls", ".xlsx"}
GENERATED_DOCUMENT_MODES = {"apostila", "trabalho", "pdf", "flashcards"}
MAX_EXTRACTED_TEXT_CHARS = 36_000
STOP_WORDS = {
    "a",
    "agora",
    "ao",
    "aos",
    "as",
    "com",
    "da",
    "das",
    "de",
    "do",
    "dos",
    "e",
    "em",
    "esse",
    "esta",
    "isso",
    "na",
    "nas",
    "no",
    "nos",
    "o",
    "os",
    "para",
    "por",
    "que",
    "um",
    "uma",
}


class OrionFilesService:
    def status(self) -> OrionFilesStatus:
        return OrionFilesStatus(
            status="ready",
            storage_backend=settings.file_storage_backend,
            storage_path=str(settings.file_storage_root),
            max_upload_bytes=settings.file_upload_max_bytes,
            allowed_extensions=sorted(settings.file_allowed_extensions),
            blocked_extensions=sorted(settings.file_blocked_extensions),
            restrictions=[
                "per-user-storage",
                "path-traversal-blocked",
                "dangerous-extensions-blocked",
                "uploaded-files-not-executed",
                "ocr-optional",
                "generated-documents-per-user",
            ],
        )

    async def save_upload(
        self,
        *,
        upload,
        user_id: str,
        category: str = "geral",
        description: str | None = None,
    ) -> OrionFileUploadResponse:
        filename = upload.filename or "arquivo"
        content = await upload.read(settings.file_upload_max_bytes + 1)
        content_type = upload.content_type or mimetypes.guess_type(filename)[0] or "application/octet-stream"
        record = self.save_bytes(
            content=content,
            user_id=user_id,
            original_name=filename,
            content_type=content_type,
            category=category,
            source="upload",
            description=description,
        )
        LOGGER.info("orion_file_uploaded user=%s file=%s size=%s", user_id, record.id, record.size_bytes)
        return OrionFileUploadResponse(
            file=record,
            message="Arquivo recebido com seguranca. Posso analisar quando voce pedir.",
        )

    def save_camera_photo(
        self,
        *,
        image_data: str,
        user_id: str,
        filename: str,
        category: str = "camera",
        description: str | None = None,
    ) -> OrionFileUploadResponse:
        content_type, content = parse_data_url(image_data)
        record = self.save_bytes(
            content=content,
            user_id=user_id,
            original_name=filename,
            content_type=content_type,
            category=category,
            source="camera",
            description=description,
        )
        LOGGER.info("orion_camera_photo_saved user=%s file=%s size=%s", user_id, record.id, record.size_bytes)
        return OrionFileUploadResponse(
            file=record,
            message="Foto salva. Posso observar os detalhes basicos agora.",
        )

    def save_bytes(
        self,
        *,
        content: bytes,
        user_id: str,
        original_name: str,
        content_type: str,
        category: str,
        source: str,
        description: str | None = None,
    ) -> OrionFileRecord:
        validate_user_id(user_id)
        validate_size(content)
        extension = validate_extension(original_name, content_type)
        file_id = uuid4().hex
        safe_name = secure_filename(original_name, extension)
        relative_path = Path("users") / user_id / f"{file_id}_{safe_name}"
        target_path = resolve_storage_path(relative_path)
        target_path.parent.mkdir(parents=True, exist_ok=True)
        target_path.write_bytes(content)

        record = repositories.create_file_record(
            file_id=file_id,
            user_id=user_id,
            original_name=Path(original_name).name or safe_name,
            safe_name=safe_name,
            content_type=content_type,
            extension=extension,
            size_bytes=len(content),
            category=safe_category(category),
            source=source,
            storage_path=relative_path.as_posix(),
            description=description,
        )
        return file_record_from_row(record)

    def list_files(self, *, user_id: str, limit: int = 50) -> list[OrionFileRecord]:
        validate_user_id(user_id)
        rows = repositories.list_file_records(user_id, limit=min(max(limit, 1), 100))
        return [file_record_from_row(row) for row in rows]

    def get_file(self, *, file_id: str, user_id: str) -> OrionFileRecord:
        validate_user_id(user_id)
        row = repositories.get_file_record(file_id, user_id)
        if row is None:
            raise FileNotFoundError("Arquivo nao encontrado para este usuario.")
        return file_record_from_row(row)

    def get_download(self, *, file_id: str, user_id: str) -> tuple[OrionFileRecord, Path]:
        validate_user_id(user_id)
        row = repositories.get_file_record(file_id, user_id)
        if row is None:
            raise FileNotFoundError("Arquivo nao encontrado para este usuario.")
        file_path = resolve_storage_path(Path(row["storage_path"]))
        if not file_path.exists():
            raise FileNotFoundError("Arquivo fisico nao encontrado no armazenamento local.")
        return file_record_from_row(row), file_path

    def delete_file(self, *, file_id: str, user_id: str) -> OrionFileRecord:
        validate_user_id(user_id)
        row = repositories.delete_file_record(file_id, user_id)
        if row is None:
            raise FileNotFoundError("Arquivo nao encontrado para este usuario.")
        file_path = resolve_storage_path(Path(row["storage_path"]))
        try:
            file_path.unlink(missing_ok=True)
        except OSError:
            LOGGER.warning("orion_file_delete_failed user=%s file=%s path=%s", user_id, file_id, file_path)
        return file_record_from_row(row)

    def analyze_file(
        self,
        *,
        file_id: str,
        user_id: str,
        instructions: str | None = None,
        manual_description: str | None = None,
    ) -> OrionFileAnalysisResponse:
        validate_user_id(user_id)
        row = repositories.get_file_record(file_id, user_id)
        if row is None:
            raise FileNotFoundError("Arquivo nao encontrado para este usuario.")

        file_path = resolve_storage_path(Path(row["storage_path"]))
        if not file_path.exists():
            summary = "O registro existe, mas o arquivo fisico nao foi encontrado no armazenamento local."
            updated = repositories.update_file_analysis(
                file_id=file_id,
                user_id=user_id,
                analysis_status="error",
                summary=summary,
                keywords_json="[]",
                description=manual_description,
            )
            file_record = file_record_from_row(updated or row)
            return OrionFileAnalysisResponse(
                file=file_record,
                summary=summary,
                keywords=[],
                message="Nao consegui abrir o arquivo salvo.",
            )

        content = file_path.read_bytes()
        analysis = analyze_content(
            content=content,
            filename=row["original_name"],
            extension=row["extension"],
            content_type=row["content_type"],
            manual_description=manual_description,
            instructions=instructions,
        )
        updated = repositories.update_file_analysis(
            file_id=file_id,
            user_id=user_id,
            analysis_status=analysis["status"],
            summary=analysis["summary"],
            keywords_json=json.dumps(analysis["keywords"], ensure_ascii=False),
            description=manual_description,
        )
        repositories.upsert_user_summary(
            user_id,
            f"Arquivo {row['original_name']}: {analysis['summary'][:180]}",
            source_type="file",
        )
        for keyword in analysis["keywords"][:5]:
            repositories.upsert_user_memory_fact(user_id, "topic", keyword)

        file_record = file_record_from_row(updated or row)
        LOGGER.info("orion_file_analyzed user=%s file=%s status=%s", user_id, file_id, analysis["status"])
        return OrionFileAnalysisResponse(
            file=file_record,
            summary=analysis["summary"],
            keywords=analysis["keywords"],
            message=build_analysis_message(analysis["status"], analysis["summary"]),
        )

    def transform_file(
        self,
        *,
        file_id: str,
        user_id: str,
        mode: Literal["summary", "explanation", "apostila", "trabalho", "pdf", "flashcards"],
        output_format: Literal["text", "pdf"] = "text",
        instructions: str | None = None,
    ) -> OrionFileTransformResponse:
        validate_user_id(user_id)
        row = repositories.get_file_record(file_id, user_id)
        if row is None:
            raise FileNotFoundError("Arquivo nao encontrado para este usuario.")

        file_path = resolve_storage_path(Path(row["storage_path"]))
        if not file_path.exists():
            raise FileNotFoundError("Arquivo fisico nao encontrado no armazenamento local.")

        content = file_path.read_bytes()
        extracted_text = extract_readable_text(
            content=content,
            filename=row["original_name"],
            extension=row["extension"],
            content_type=row["content_type"],
        )
        if not extracted_text:
            analysis = analyze_content(
                content=content,
                filename=row["original_name"],
                extension=row["extension"],
                content_type=row["content_type"],
                manual_description=row.get("description"),
                instructions=instructions,
            )
            extracted_text = analysis["summary"]

        title = transform_title(mode, row["original_name"])
        transformed_content = build_transformed_content(
            text=extracted_text,
            filename=row["original_name"],
            mode=mode,
            instructions=instructions,
        )
        target_format = "pdf" if mode == "pdf" or output_format == "pdf" else "text"
        generated_file = None

        if mode in GENERATED_DOCUMENT_MODES or target_format == "pdf":
            generated_file = self._save_generated_document(
                user_id=user_id,
                source_row=row,
                title=title,
                content=transformed_content,
                mode=mode,
                output_format=target_format,
            )

        LOGGER.info(
            "orion_file_transformed user=%s source=%s mode=%s format=%s generated=%s",
            user_id,
            file_id,
            mode,
            target_format,
            generated_file.id if generated_file else None,
        )
        return OrionFileTransformResponse(
            source_file=file_record_from_row(row),
            generated_file=generated_file,
            mode=mode,
            output_format=target_format,
            title=title,
            content=transformed_content,
            message=build_transform_message(mode, target_format, generated_file),
        )

    def _save_generated_document(
        self,
        *,
        user_id: str,
        source_row: dict,
        title: str,
        content: str,
        mode: str,
        output_format: Literal["text", "pdf"],
    ) -> OrionFileRecord:
        source_stem = Path(source_row["safe_name"]).stem or "documento"
        if output_format == "pdf":
            generated_bytes = generate_pdf_document(title=title, content=content)
            extension = ".pdf"
            content_type = "application/pdf"
        else:
            generated_bytes = content.encode("utf-8")
            extension = ".txt"
            content_type = "text/plain; charset=utf-8"

        record = self.save_bytes(
            content=generated_bytes,
            user_id=user_id,
            original_name=f"{source_stem}-{mode}{extension}",
            content_type=content_type,
            category="gerado",
            source="generated",
            description=f"Gerado pelo Orion a partir de {source_row['original_name']}.",
        )
        summary = f"{title} gerado a partir de {source_row['original_name']}."
        updated = repositories.update_file_analysis(
            file_id=record.id,
            user_id=user_id,
            analysis_status="ready",
            summary=summary,
            keywords_json=json.dumps(extract_keywords(content), ensure_ascii=False),
            description=record.description,
        )
        return file_record_from_row(updated or record.model_dump())


def validate_user_id(user_id: str) -> None:
    if not USER_ID_PATTERN.fullmatch(user_id):
        raise ValueError("Identificador de usuario invalido.")


def validate_size(content: bytes) -> None:
    if not content:
        raise ValueError("Arquivo vazio nao pode ser salvo.")
    if len(content) > settings.file_upload_max_bytes:
        raise ValueError("Arquivo excede o limite de tamanho permitido.")


def validate_extension(original_name: str, content_type: str) -> str:
    guessed_extension = mimetypes.guess_extension(content_type.split(";")[0].strip()) or ""
    extension = Path(original_name).suffix.lower() or guessed_extension.lower()
    if extension == ".jpe":
        extension = ".jpg"
    if extension in settings.file_blocked_extensions:
        raise ValueError("Este tipo de arquivo e bloqueado por seguranca.")
    if extension not in settings.file_allowed_extensions:
        raise ValueError("Tipo de arquivo nao permitido nesta instalacao.")
    return extension


def safe_category(category: str) -> str:
    normalized = normalize_ascii(category or "geral").lower()
    clean = SAFE_NAME_PATTERN.sub("-", normalized).strip(".-_")
    return clean[:40] or "geral"


def secure_filename(original_name: str, extension: str) -> str:
    name = Path(original_name).name
    stem = Path(name).stem or "arquivo"
    normalized = normalize_ascii(stem)
    clean_stem = SAFE_NAME_PATTERN.sub("-", normalized).strip(".-_")[:72] or "arquivo"
    return f"{clean_stem}{extension}"


def normalize_ascii(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    return normalized.encode("ascii", "ignore").decode("ascii")


def resolve_storage_path(relative_path: Path) -> Path:
    root = settings.file_storage_root.resolve()
    target = (root / relative_path).resolve()
    if not target.is_relative_to(root):
        raise ValueError("Caminho de arquivo invalido.")
    return target


def parse_data_url(image_data: str) -> tuple[str, bytes]:
    if not image_data.startswith("data:"):
        raise ValueError("Foto deve ser enviada como data URL.")
    header, _, encoded = image_data.partition(",")
    if ";base64" not in header or not encoded:
        raise ValueError("Foto deve estar codificada em base64.")
    content_type = header.removeprefix("data:").split(";", 1)[0] or "image/png"
    if content_type not in {"image/png", "image/jpeg", "image/webp"}:
        raise ValueError("Formato de foto nao permitido.")
    try:
        return content_type, base64.b64decode(encoded, validate=True)
    except (binascii.Error, ValueError) as exc:
        raise ValueError("Foto enviada esta corrompida ou invalida.") from exc


def file_record_from_row(row: dict) -> OrionFileRecord:
    keywords = []
    try:
        keywords = json.loads(row.get("keywords_json") or "[]")
    except json.JSONDecodeError:
        keywords = []
    return OrionFileRecord(
        id=row["id"],
        user_id=row["user_id"],
        original_name=row["original_name"],
        safe_name=row["safe_name"],
        content_type=row["content_type"],
        extension=row["extension"],
        size_bytes=row["size_bytes"],
        category=row["category"],
        source=row["source"],
        analysis_status=row["analysis_status"],
        summary=row.get("summary"),
        keywords=[str(keyword) for keyword in keywords if str(keyword).strip()],
        description=row.get("description"),
        created_at=row["created_at"],
        updated_at=row["updated_at"],
    )


def analyze_content(
    *,
    content: bytes,
    filename: str,
    extension: str,
    content_type: str,
    manual_description: str | None,
    instructions: str | None,
) -> dict:
    if extension in TEXT_EXTENSIONS | {".docx", ".pdf", ".pptx", ".xlsx"}:
        text = extract_readable_text(
            content=content,
            filename=filename,
            extension=extension,
            content_type=content_type,
        )
        if not text and extension == ".pdf":
            summary = (
                "PDF salvo com seguranca. Nao encontrei texto extraivel automaticamente; "
                "PDF escaneado precisa de OCR configurado."
            )
            return {
                "status": "limited",
                "summary": append_manual_description(summary, manual_description),
                "keywords": extract_keywords(filename),
            }
        if not text:
            summary = (
                f"Documento {extension} salvo com seguranca. Nao encontrei texto legivel para extrair automaticamente."
            )
            return {
                "status": "limited",
                "summary": append_manual_description(summary, manual_description),
                "keywords": extract_keywords(filename),
            }
        summary = summarize_text(text, filename=filename, instructions=instructions)
        keywords = extract_keywords(text)
        return {
            "status": "ready",
            "summary": append_manual_description(summary, manual_description),
            "keywords": keywords,
        }
    if extension in IMAGE_EXTENSIONS:
        image_info = describe_image(content, content_type)
        base = f"Imagem {image_info} salva com {len(content)} bytes"
        summary = (
            f"{base}. Consigo guardar, pre-visualizar e receber uma descricao manual. "
            "Leitura automatica de texto em imagem precisa de OCR configurado."
        )
        return {
            "status": "limited",
            "summary": append_manual_description(summary, manual_description),
            "keywords": ["imagem", "camera" if "camera" in filename.lower() else "arquivo"],
        }
    if extension in DOCUMENT_EXTENSIONS:
        summary = (
            f"Documento {extension} salvo com seguranca. Para leitura automatica completa, "
            "configure um extrator compativel ou envie uma versao em PDF/texto."
        )
        return {
            "status": "limited",
            "summary": append_manual_description(summary, manual_description),
            "keywords": extract_keywords(filename),
        }
    summary = "Arquivo salvo e catalogado. Analise automatica completa ainda nao esta disponivel para este tipo."
    return {
        "status": "limited",
        "summary": append_manual_description(summary, manual_description),
        "keywords": extract_keywords(filename),
    }


def decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def extract_readable_text(
    *,
    content: bytes,
    filename: str,
    extension: str,
    content_type: str,
) -> str:
    if extension in TEXT_EXTENSIONS:
        return limit_extracted_text(decode_text(content))
    if extension == ".pdf":
        return limit_extracted_text(extract_pdf_text(content))
    if extension == ".docx":
        return limit_extracted_text(extract_docx_text(content))
    if extension == ".xlsx":
        return limit_extracted_text(extract_xlsx_text(content))
    if extension == ".pptx":
        return limit_extracted_text(extract_pptx_text(content))
    if extension in IMAGE_EXTENSIONS:
        return describe_image(content, content_type)
    return ""


def extract_pdf_text(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception:
        return ""
    try:
        reader = PdfReader(BytesIO(content))
        parts = [page.extract_text() or "" for page in reader.pages[:20]]
    except Exception:
        return ""
    return "\n".join(parts).strip()


def extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except Exception:
        return ""
    try:
        document = Document(BytesIO(content))
    except Exception:
        return ""

    parts: list[str] = []
    parts.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables[:12]:
        for row in table.rows[:80]:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts).strip()


def extract_xlsx_text(content: bytes) -> str:
    try:
        with zipfile.ZipFile(BytesIO(content)) as archive:
            shared_strings = read_xlsx_shared_strings(archive)
            sheet_names = sorted(
                name
                for name in archive.namelist()
                if name.startswith("xl/worksheets/") and name.endswith(".xml")
            )[:12]
            parts: list[str] = []
            for sheet_index, sheet_name in enumerate(sheet_names, start=1):
                xml_text = archive.read(sheet_name)
                rows = read_xlsx_sheet_rows(xml_text, shared_strings)
                if rows:
                    parts.append(f"Planilha {sheet_index}")
                    parts.extend(rows[:120])
            return "\n".join(parts).strip()
    except Exception:
        return ""


def read_xlsx_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    try:
        xml_text = archive.read("xl/sharedStrings.xml")
    except KeyError:
        return []
    try:
        root = ElementTree.fromstring(xml_text)
    except ElementTree.ParseError:
        return []
    strings: list[str] = []
    for item in root.iter():
        if item.tag.endswith("}si") or item.tag == "si":
            values = [
                text_node.text or ""
                for text_node in item.iter()
                if text_node.tag.endswith("}t") or text_node.tag == "t"
            ]
            strings.append("".join(values).strip())
    return strings


def read_xlsx_sheet_rows(xml_text: bytes, shared_strings: list[str]) -> list[str]:
    try:
        root = ElementTree.fromstring(xml_text)
    except ElementTree.ParseError:
        return []
    rows: list[str] = []
    for row in root.iter():
        if not (row.tag.endswith("}row") or row.tag == "row"):
            continue
        values = [read_xlsx_cell(cell, shared_strings) for cell in row if cell.tag.endswith("}c") or cell.tag == "c"]
        values = [value for value in values if value]
        if values:
            rows.append(" | ".join(values))
        if len(rows) >= 120:
            break
    return rows


def read_xlsx_cell(cell: ElementTree.Element, shared_strings: list[str]) -> str:
    cell_type = cell.attrib.get("t", "")
    if cell_type == "inlineStr":
        values = [
            text_node.text or ""
            for text_node in cell.iter()
            if text_node.tag.endswith("}t") or text_node.tag == "t"
        ]
        return "".join(values).strip()

    value_node = next(
        (node for node in cell if node.tag.endswith("}v") or node.tag == "v"),
        None,
    )
    if value_node is None or value_node.text is None:
        return ""
    raw_value = value_node.text.strip()
    if cell_type == "s":
        try:
            return shared_strings[int(raw_value)]
        except (IndexError, ValueError):
            return raw_value
    return raw_value


def extract_pptx_text(content: bytes) -> str:
    try:
        with zipfile.ZipFile(BytesIO(content)) as archive:
            slide_names = sorted(
                name
                for name in archive.namelist()
                if name.startswith("ppt/slides/slide") and name.endswith(".xml")
            )[:80]
            parts: list[str] = []
            for slide_index, slide_name in enumerate(slide_names, start=1):
                slide_text = read_pptx_slide_text(archive.read(slide_name))
                if slide_text:
                    parts.append(f"Slide {slide_index}: {slide_text}")
            return "\n".join(parts).strip()
    except Exception:
        return ""


def read_pptx_slide_text(xml_text: bytes) -> str:
    try:
        root = ElementTree.fromstring(xml_text)
    except ElementTree.ParseError:
        return ""
    values = [
        node.text.strip()
        for node in root.iter()
        if (node.tag.endswith("}t") or node.tag == "t") and node.text and node.text.strip()
    ]
    return " ".join(values)


def limit_extracted_text(text: str) -> str:
    clean = text.strip()
    if len(clean) <= MAX_EXTRACTED_TEXT_CHARS:
        return clean
    return clean[:MAX_EXTRACTED_TEXT_CHARS].rsplit(" ", 1)[0].strip()


def summarize_text(text: str, *, filename: str, instructions: str | None) -> str:
    clean = " ".join(text.split())
    if not clean:
        return f"O arquivo {filename} foi salvo, mas nao encontrei texto legivel para resumir."
    sentences = split_sentences(clean)
    chosen = sentences[:3] if sentences else [clean[:500]]
    summary = " ".join(chosen)
    if len(summary) > 700:
        summary = f"{summary[:697].rstrip()}..."
    prefix = f"Li {filename}. "
    if instructions:
        prefix += f"Pedido considerado: {instructions.strip()[:160]}. "
    return f"{prefix}Resumo: {summary}"


def build_transformed_content(
    *,
    text: str,
    filename: str,
    mode: str,
    instructions: str | None,
) -> str:
    clean = " ".join(text.split())
    summary = summarize_text(clean, filename=filename, instructions=instructions)
    keywords = extract_keywords(clean)
    sentences = split_sentences(clean)
    if not sentences and clean:
        sentences = [clean[:500]]

    if mode == "summary":
        return summary
    if mode == "explanation":
        return build_explanation(filename=filename, summary=summary, keywords=keywords, sentences=sentences)
    if mode == "apostila":
        return build_study_handout(filename=filename, summary=summary, keywords=keywords, sentences=sentences)
    if mode == "trabalho":
        return build_school_work(filename=filename, summary=summary, keywords=keywords, sentences=sentences)
    if mode == "flashcards":
        return build_flashcards(filename=filename, keywords=keywords, sentences=sentences)
    return build_pdf_source_text(filename=filename, summary=summary, keywords=keywords, sentences=sentences)


def build_explanation(*, filename: str, summary: str, keywords: list[str], sentences: list[str]) -> str:
    details = "\n".join(f"- {sentence}" for sentence in sentences[:6])
    key_text = ", ".join(keywords[:8]) or "pontos principais"
    return (
        f"Explicacao do arquivo {filename}\n\n"
        f"{summary}\n\n"
        f"Em linguagem simples: o material gira em torno de {key_text}. "
        "O melhor caminho e separar a ideia central, os detalhes de apoio e as possiveis duvidas.\n\n"
        f"Pontos explicados:\n{details}"
    )


def build_study_handout(*, filename: str, summary: str, keywords: list[str], sentences: list[str]) -> str:
    objectives = "\n".join(f"- Entender {keyword}" for keyword in keywords[:5]) or "- Entender o conteudo principal"
    topics = "\n".join(f"{index}. {sentence}" for index, sentence in enumerate(sentences[:8], start=1))
    exercises = "\n".join(
        f"{index}. Explique com suas palavras: {keyword}."
        for index, keyword in enumerate((keywords or ["tema principal"])[:5], start=1)
    )
    return (
        f"Apostila Orion - {filename}\n\n"
        "Objetivos de aprendizagem\n"
        f"{objectives}\n\n"
        "Resumo orientado\n"
        f"{summary}\n\n"
        "Conteudo organizado\n"
        f"{topics}\n\n"
        "Atividades\n"
        f"{exercises}\n\n"
        "Revisao final\n"
        "Leia o resumo, responda as atividades e marque os termos que ainda precisam de explicacao."
    )


def build_school_work(*, filename: str, summary: str, keywords: list[str], sentences: list[str]) -> str:
    development = "\n".join(f"Paragrafo {index}: {sentence}" for index, sentence in enumerate(sentences[:6], start=1))
    references = (
        "\n".join(f"- Topico derivado do arquivo: {keyword}" for keyword in keywords[:6])
        or "- Arquivo enviado"
    )
    return (
        f"Trabalho Orion - {filename}\n\n"
        "Introducao\n"
        f"Este trabalho apresenta uma organizacao do conteudo extraido de {filename}.\n\n"
        "Desenvolvimento\n"
        f"{development}\n\n"
        "Conclusao\n"
        f"{summary}\n\n"
        "Topicos para referencias e revisao\n"
        f"{references}"
    )


def build_flashcards(*, filename: str, keywords: list[str], sentences: list[str]) -> str:
    cards: list[str] = [f"Flashcards Orion - {filename}"]
    base_keywords = keywords[:8] or ["tema principal", "ideia central", "detalhe importante"]
    for index, keyword in enumerate(base_keywords, start=1):
        answer = (
            sentence_for_keyword(keyword, sentences)
            or "Revise o trecho principal do documento e explique com suas palavras."
        )
        cards.append(f"\nCartao {index}\nPergunta: O que o material diz sobre {keyword}?\nResposta: {answer}")
    return "\n".join(cards)


def build_pdf_source_text(*, filename: str, summary: str, keywords: list[str], sentences: list[str]) -> str:
    highlights = "\n".join(f"- {sentence}" for sentence in sentences[:10])
    return (
        f"Documento PDF Orion - {filename}\n\n"
        f"{summary}\n\n"
        f"Palavras-chave: {', '.join(keywords[:10]) or 'conteudo do arquivo'}\n\n"
        f"Destaques\n{highlights}"
    )


def sentence_for_keyword(keyword: str, sentences: list[str]) -> str:
    normalized_keyword = normalize_ascii(keyword).lower()
    for sentence in sentences:
        if normalized_keyword in normalize_ascii(sentence).lower():
            return sentence
    return sentences[0] if sentences else ""


def transform_title(mode: str, filename: str) -> str:
    labels = {
        "summary": "Resumo",
        "explanation": "Explicacao",
        "apostila": "Apostila",
        "trabalho": "Trabalho",
        "pdf": "PDF",
        "flashcards": "Flashcards",
    }
    return f"{labels.get(mode, 'Documento')} Orion - {filename}"


def build_transform_message(mode: str, output_format: str, generated_file: OrionFileRecord | None) -> str:
    if generated_file:
        return (
            f"Transformacao {mode} concluida em {output_format}. "
            f"Gerei o arquivo {generated_file.original_name} e deixei em Meus Arquivos."
        )
    return f"Transformacao {mode} concluida. Posso gerar PDF se voce quiser."


def generate_pdf_document(*, title: str, content: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except Exception as exc:
        raise ValueError("Geracao de PDF indisponivel. Instale reportlab para ativar esta funcao.") from exc

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        title=title,
        leftMargin=48,
        rightMargin=48,
        topMargin=54,
        bottomMargin=54,
    )
    styles = getSampleStyleSheet()
    story = [Paragraph(escape_pdf_text(title), styles["Title"]), Spacer(1, 14)]
    for block in split_pdf_blocks(content):
        style_name = "Heading2" if len(block) < 90 and not block.endswith(".") else "BodyText"
        story.append(Paragraph(escape_pdf_text(block), styles[style_name]))
        story.append(Spacer(1, 8))
    document.build(story)
    return buffer.getvalue()


def split_pdf_blocks(content: str) -> list[str]:
    blocks = [block.strip() for block in re.split(r"\n{2,}", content) if block.strip()]
    expanded: list[str] = []
    for block in blocks:
        if len(block) <= 900:
            expanded.append(block)
            continue
        sentences = split_sentences(block)
        expanded.extend(sentences or [block[:900]])
    return expanded[:80]


def escape_pdf_text(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def split_sentences(text: str) -> list[str]:
    raw_sentences = re.split(r"(?<=[.!?])\s+", text)
    return [sentence.strip() for sentence in raw_sentences if len(sentence.strip()) > 20]


def extract_keywords(text: str) -> list[str]:
    normalized = unicodedata.normalize("NFKD", text.lower())
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    tokens = re.findall(r"[a-z0-9]{3,}", ascii_text)
    counter = Counter(token for token in tokens if token not in STOP_WORDS)
    return [word for word, _count in counter.most_common(10)]


def append_manual_description(summary: str, manual_description: str | None) -> str:
    if not manual_description:
        return summary
    return f"{summary} Descricao manual do usuario: {manual_description.strip()[:500]}"


def describe_image(content: bytes, content_type: str) -> str:
    try:
        from PIL import Image
    except Exception:
        dimensions = detect_image_dimensions(content)
        if dimensions:
            return f"{content_type} com dimensoes aproximadas de {dimensions[0]}x{dimensions[1]} pixels"
        return content_type

    try:
        with Image.open(BytesIO(content)) as image:
            width, height = image.size
            image_format = image.format or content_type
            mode = image.mode
            return f"{image_format} {width}x{height} pixels, modo {mode}"
    except Exception:
        dimensions = detect_image_dimensions(content)
        if dimensions:
            return f"{content_type} com dimensoes aproximadas de {dimensions[0]}x{dimensions[1]} pixels"
        return content_type


def detect_image_dimensions(content: bytes) -> tuple[int, int] | None:
    if content.startswith(b"\x89PNG\r\n\x1a\n") and len(content) >= 24:
        return int.from_bytes(content[16:20], "big"), int.from_bytes(content[20:24], "big")
    if content[:6] in {b"GIF87a", b"GIF89a"} and len(content) >= 10:
        return int.from_bytes(content[6:8], "little"), int.from_bytes(content[8:10], "little")
    if content.startswith(b"\xff\xd8"):
        return detect_jpeg_dimensions(content)
    return None


def detect_jpeg_dimensions(content: bytes) -> tuple[int, int] | None:
    index = 2
    while index + 9 < len(content):
        if content[index] != 0xFF:
            index += 1
            continue
        marker = content[index + 1]
        length = int.from_bytes(content[index + 2 : index + 4], "big")
        if marker in {0xC0, 0xC1, 0xC2, 0xC3, 0xC5, 0xC6, 0xC7, 0xC9, 0xCA, 0xCB, 0xCD, 0xCE, 0xCF}:
            height = int.from_bytes(content[index + 5 : index + 7], "big")
            width = int.from_bytes(content[index + 7 : index + 9], "big")
            return width, height
        if length < 2:
            break
        index += 2 + length
    return None


def build_analysis_message(status: str, summary: str) -> str:
    if status == "ready":
        return f"Li o arquivo. {summary} Posso explicar, criar um resumo de estudo ou pesquisar fontes atuais."
    if status == "limited":
        return f"Analise basica concluida. {summary} Posso organizar isso e pesquisar referencias se voce autorizar."
    return summary
